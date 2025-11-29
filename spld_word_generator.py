"""SPLD Word Document Generator - Exact Format"""

from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_color(cell, color_hex):
    """Set background color of a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)


def create_spld_word_report(projects, output_path):
    """Generate SPLD format Word document FROM EXCEL DATA."""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(8.5)
        section.page_width = Inches(11)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.orientation = 1  # Landscape
    
    for i, project in enumerate(projects):
        create_spld_project_page_word(doc, project)
        
        # Add page break except for last project
        if i < len(projects) - 1:
            doc.add_page_break()
    
    doc.save(output_path)
    return output_path


def create_spld_project_page_word(doc, project):
    """Create SPLD format project page in Word FROM EXCEL DATA."""
    
    # Header Table
    header_table = doc.add_table(rows=3, cols=4)
    header_table.style = 'Table Grid'
    
    # Row 1
    header_table.cell(0, 0).text = 'Project Status Report'
    header_table.cell(0, 2).text = 'Project Sponsor'
    header_table.cell(0, 3).text = str(project.get('gm', 'TBD'))
    
    # Row 2 - Project Name
    header_table.cell(1, 0).text = str(project.get('name', 'Unnamed Project'))
    header_table.cell(1, 2).text = 'Project Manager'
    header_table.cell(1, 3).text = str(project.get('operational_lead', 'TBD'))
    
    # Row 3 - Report Date
    header_table.cell(2, 0).text = f'Report Date: {datetime.now().strftime("%d/%m/%Y")}'
    
    # Merge cells
    header_table.cell(0, 0).merge(header_table.cell(0, 1))
    header_table.cell(1, 0).merge(header_table.cell(1, 1))
    header_table.cell(2, 0).merge(header_table.cell(2, 1))
    
    # Style header
    for row in header_table.rows:
        for cell in row.cells:
            set_cell_color(cell, '2C3E50')  # Gray background
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Orange project name
    header_table.cell(1, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(234, 106, 31)
    header_table.cell(1, 0).paragraphs[0].runs[0].font.bold = True
    header_table.cell(1, 0).paragraphs[0].runs[0].font.size = Pt(14)
    
    doc.add_paragraph()  # Spacer
    
    # Two Column Layout
    layout_table = doc.add_table(rows=1, cols=2)
    left_cell = layout_table.rows[0].cells[0]
    right_cell = layout_table.rows[0].cells[1]
    
    # LEFT COLUMN - Dates and Objective
    
    # Dates Table
    end_date_str = str(project.get('contract_end_date', 'TBD'))
    dates_table = left_cell.add_table(rows=3, cols=2)
    dates_table.style = 'Table Grid'
    
    dates_data = [
        ['Start Date', datetime.now().strftime('%d %b %Y')],
        ['Baseline Finish', end_date_str],
        ['Forecast Finish', end_date_str]
    ]
    
    for i, row_data in enumerate(dates_data):
        dates_table.cell(i, 0).text = row_data[0]
        dates_table.cell(i, 1).text = row_data[1]
        set_cell_color(dates_table.cell(i, 0), '2C3E50')
        dates_table.cell(i, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    left_cell.add_paragraph()
    
    # Project Objective
    obj_heading = left_cell.add_paragraph('Project Objective')
    obj_heading.runs[0].font.bold = True
    obj_heading.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    obj_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Get objective text from Excel
    objective_text = (project.get('description', '') or 
                     project.get('comments', '') or 
                     f"Managing and delivering {project.get('name', 'project')} successfully")[:300]
    
    obj_para = left_cell.add_paragraph(objective_text)
    obj_para.runs[0].font.size = Pt(9)
    
    left_cell.add_paragraph()
    
    # Deliverables Table
    deliv_heading = left_cell.add_paragraph('Deliverables')
    deliv_heading.runs[0].font.bold = True
    deliv_heading.runs[0].font.color.rgb = RGBColor(234, 106, 31)
    deliv_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create deliverables based on actual progress
    progress = project.get('timeline_actual', 0)
    end_date = project.get('contract_end_date', 'TBD')
    
    deliv_table = left_cell.add_table(rows=3, cols=5)
    deliv_table.style = 'Table Grid'
    
    # Headers
    headers = ['Deliverable', 'Contract', 'Planned', '% Done', 'Status']
    for i, header in enumerate(headers):
        deliv_table.cell(0, i).text = header
        set_cell_color(deliv_table.cell(0, i), '2C3E50')
        deliv_table.cell(0, i).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        deliv_table.cell(0, i).paragraphs[0].runs[0].font.bold = True
    
    # Add deliverable rows based on progress
    if progress >= 75:
        deliv_table.cell(1, 0).text = 'Development'
        deliv_table.cell(1, 3).text = f'{progress:.0f}%'
        deliv_table.cell(1, 4).text = 'On Track'
        set_cell_color(deliv_table.cell(1, 4), '27AE60')
        
        deliv_table.cell(2, 0).text = 'Testing'
        deliv_table.cell(2, 3).text = f'{max(0, progress-80):.0f}%'
        deliv_table.cell(2, 4).text = 'In Progress'
    else:
        deliv_table.cell(1, 0).text = 'Planning'
        deliv_table.cell(1, 3).text = f'{min(100, progress*2):.0f}%'
        deliv_table.cell(1, 4).text = 'On Track' if progress > 0 else 'Not Started'
        
        deliv_table.cell(2, 0).text = 'Implementation'
        deliv_table.cell(2, 3).text = f'{progress:.0f}%'
        deliv_table.cell(2, 4).text = 'In Progress' if progress > 25 else 'Not Started'
    
    # Fill dates
    for i in range(1, 3):
        deliv_table.cell(i, 1).text = str(end_date)
        deliv_table.cell(i, 2).text = str(end_date)
    
    # RIGHT COLUMN - Health and Activities
    
    # Overall Health
    health_heading = right_cell.add_paragraph('Overall Project Health')
    health_heading.runs[0].font.bold = True
    health_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Determine health status from Excel data
    health = str(project.get('health', 'On Track'))
    if 'on track' in health.lower():
        health_text = 'On Track'
        health_color = RGBColor(39, 174, 96)
    elif 'at risk' in health.lower():
        health_text = 'Slightly Delayed'
        health_color = RGBColor(243, 156, 18)
    else:
        health_text = 'Delayed'
        health_color = RGBColor(231, 76, 60)
    
    health_status = right_cell.add_paragraph(health_text)
    health_status.runs[0].font.bold = True
    health_status.runs[0].font.size = Pt(12)
    health_status.runs[0].font.color.rgb = health_color
    health_status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    right_cell.add_paragraph()
    
    # Project Progress
    progress_heading = right_cell.add_paragraph('Project Progress')
    progress_heading.runs[0].font.bold = True
    progress_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    actual_progress = project.get('timeline_actual', 0)
    planned_progress = project.get('timeline_planned', 0)
    
    progress_text = right_cell.add_paragraph(f'{actual_progress:.0f}%')
    progress_text.runs[0].font.size = Pt(16)
    progress_text.runs[0].font.bold = True
    progress_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    planned_text = right_cell.add_paragraph(f'Planned: {planned_progress:.0f}%')
    planned_text.runs[0].font.size = Pt(9)
    planned_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    right_cell.add_paragraph()
    
    # Activity Progress
    activity_heading = right_cell.add_paragraph('Activity Progress')
    activity_heading.runs[0].font.bold = True
    activity_heading.runs[0].font.color.rgb = RGBColor(234, 106, 31)
    activity_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Current Activities from Excel
    current_heading = right_cell.add_paragraph('Current Activities')
    current_heading.runs[0].font.bold = True
    current_heading.runs[0].font.size = Pt(10)
    
    current_activities = str(project.get('current_activities', '[To be provided]'))
    if len(current_activities) > 200:
        current_activities = current_activities[:197] + '...'
    
    current_text = right_cell.add_paragraph(current_activities)
    current_text.runs[0].font.size = Pt(9)
    
    # Future Activities from Excel
    future_heading = right_cell.add_paragraph('Future Activities')
    future_heading.runs[0].font.bold = True
    future_heading.runs[0].font.size = Pt(10)
    
    future_activities = str(project.get('future_activities', '[To be provided]'))
    if len(future_activities) > 200:
        future_activities = future_activities[:197] + '...'
    
    future_text = right_cell.add_paragraph(future_activities)
    future_text.runs[0].font.size = Pt(9)
    
    doc.add_paragraph()  # Spacer
    
    # Risks and Issues Section
    risks_heading = doc.add_paragraph('Risks & Issues')
    risks_heading.runs[0].font.bold = True
    risks_heading.runs[0].font.color.rgb = RGBColor(234, 106, 31)
    risks_heading.runs[0].font.size = Pt(12)
    
    risks_table = doc.add_table(rows=3, cols=6)
    risks_table.style = 'Table Grid'
    
    # Headers
    headers = ['', 'Description', 'Impact', 'Mitigation', 'Owner', 'Due Date']
    for i, header in enumerate(headers):
        risks_table.cell(0, i).text = header
        set_cell_color(risks_table.cell(0, i), '2C3E50')
        risks_table.cell(0, i).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        risks_table.cell(0, i).paragraphs[0].runs[0].font.bold = True
    
    # Add risks and issues from Excel data
    risks_table.cell(1, 0).text = 'Risk'
    risks_table.cell(1, 1).text = str(project.get('risks', '[To be provided]'))[:100]
    risks_table.cell(1, 2).text = 'Medium'
    risks_table.cell(1, 3).text = 'Under review'
    risks_table.cell(1, 4).text = str(project.get('operational_lead', 'TBD'))
    risks_table.cell(1, 5).text = str(project.get('contract_end_date', 'TBD'))
    
    risks_table.cell(2, 0).text = 'Issue'
    risks_table.cell(2, 1).text = str(project.get('issues', '[To be provided]'))[:100]
    risks_table.cell(2, 2).text = 'High'
    risks_table.cell(2, 3).text = 'Being addressed'
    risks_table.cell(2, 4).text = str(project.get('operational_lead', 'TBD'))
    risks_table.cell(2, 5).text = 'ASAP'
    
    # Color impact cells
    set_cell_color(risks_table.cell(1, 2), 'F39C12')  # Yellow for Medium
    set_cell_color(risks_table.cell(2, 2), 'E74C3C')  # Red for High
    risks_table.cell(2, 2).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)


def generate_individual_spld_word(project, output_path):
    """Generate SPLD Word document for a single project."""
    return create_spld_word_report([project], output_path)
