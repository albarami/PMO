"""Word Document Generation Module for PMO Reports"""

from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_color(cell, color_hex):
    """Set background color of a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)


def create_word_report(projects, output_path):
    """Generate a Word document report for all projects."""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(8.5)
        section.page_width = Inches(11)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.orientation = 1  # Landscape
    
    for i, project in enumerate(projects):
        create_project_page_word(doc, project)
        
        # Add page break except for last project
        if i < len(projects) - 1:
            doc.add_page_break()
    
    doc.save(output_path)
    return output_path


def create_project_page_word(doc, project):
    """Create a single project page in Word document."""
    # Header
    header = doc.add_heading('Project Status Report', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Report date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f"Report Date: {datetime.now().strftime('%d/%m/%Y')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(102, 102, 102)
    
    # Project name
    project_heading = doc.add_heading(project['name'], 1)
    project_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    project_heading.runs[0].font.color.rgb = RGBColor(224, 112, 32)  # Orange
    
    # Basic info paragraph
    info_para = doc.add_paragraph()
    info_para.add_run(f"Category: {project['category']} | ")
    info_para.add_run(f"Status: {project['status']} | ")
    info_para.add_run(f"Vendor: {str(project['vendor'])[:50]}")
    
    # Key metrics table
    metrics_table = doc.add_table(rows=2, cols=5)
    metrics_table.style = 'Table Grid'
    
    # Header row
    headers = ['Sponsor GM', 'Director', 'Project Lead', 'Contract End', 'Days Remaining']
    for i, header in enumerate(headers):
        cell = metrics_table.rows[0].cells[i]
        cell.text = header
        set_cell_color(cell, '262626')  # Dark gray
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data row
    values = [
        str(project.get('gm', 'TBD')),
        str(project.get('director', 'TBD')),
        str(project.get('operational_lead', 'TBD')),
        str(project.get('contract_end_date', '[TBD]')),
        str(project.get('days_remaining', 0))
    ]
    
    for i, value in enumerate(values):
        cell = metrics_table.rows[1].cells[i]
        cell.text = value
        set_cell_color(cell, '404040')  # Medium gray
        run = cell.paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacer
    
    # Two column layout - using tables
    layout_table = doc.add_table(rows=1, cols=2)
    left_cell = layout_table.rows[0].cells[0]
    right_cell = layout_table.rows[0].cells[1]
    
    # Left column content
    left_doc = left_cell.add_paragraph()
    left_doc.add_run('Project Timeline').bold = True
    left_doc.runs[0].font.color.rgb = RGBColor(224, 112, 32)
    
    left_cell.add_paragraph(f"Actual Progress: {project['timeline_actual']:.1f}%")
    left_cell.add_paragraph(f"Planned Progress: {project['timeline_planned']:.1f}%")
    
    variance = project['schedule_variance']
    variance_text = f"Schedule Variance: {variance:+.1f}%"
    variance_para = left_cell.add_paragraph(variance_text)
    if variance >= 0:
        variance_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
    else:
        variance_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red
    
    left_cell.add_paragraph()
    
    budget_heading = left_cell.add_paragraph()
    budget_heading.add_run('Budget Utilization').bold = True
    budget_heading.runs[0].font.color.rgb = RGBColor(224, 112, 32)
    
    left_cell.add_paragraph(f"Total Budget: {project['budget_total']:,.2f} SAR")
    left_cell.add_paragraph(f"Spent: {project['budget_spent']:,.2f} SAR ({project['budget_spent_pct']:.1f}%)")
    left_cell.add_paragraph(f"Remaining: {project['budget_remaining']:,.2f} SAR ({project['budget_remaining_pct']:.1f}%)")
    
    left_cell.add_paragraph()
    
    kpi_heading = left_cell.add_paragraph()
    kpi_heading.add_run('Service Delivery KPI').bold = True
    kpi_heading.runs[0].font.color.rgb = RGBColor(224, 112, 32)
    left_cell.add_paragraph(str(project['kpi']))
    
    # Right column content
    right_doc = right_cell.add_paragraph()
    right_doc.add_run('Overall Project Health').bold = True
    right_doc.runs[0].font.color.rgb = RGBColor(224, 112, 32)
    
    # Health status
    health_text = str(project['health']) if project['health'] else 'TBD'
    health_para = right_cell.add_paragraph()
    health_run = health_para.add_run(f"  {health_text.upper()}  ")
    health_run.font.bold = True
    
    if 'on track' in health_text.lower():
        health_run.font.color.rgb = RGBColor(255, 255, 255)
        # Note: Can't easily set background color in runs, using highlight instead
        health_run.font.highlight_color = 3  # Green highlight
    elif 'at risk' in health_text.lower():
        health_run.font.color.rgb = RGBColor(255, 255, 255)
        health_run.font.highlight_color = 7  # Yellow/Orange highlight
    else:
        health_run.font.color.rgb = RGBColor(255, 255, 255)
        health_run.font.highlight_color = 6  # Red highlight
    
    right_cell.add_paragraph()
    
    activities_heading = right_cell.add_paragraph()
    activities_heading.add_run('Current Activities').bold = True
    activities_heading.runs[0].font.color.rgb = RGBColor(224, 112, 32)
    right_cell.add_paragraph(str(project['current_activities']))
    
    right_cell.add_paragraph()
    
    future_heading = right_cell.add_paragraph()
    future_heading.add_run('Future Activities').bold = True
    future_heading.runs[0].font.color.rgb = RGBColor(224, 112, 32)
    right_cell.add_paragraph(str(project['future_activities']))
    
    doc.add_paragraph()  # Spacer
    
    # Risks and Issues section
    risks_heading = doc.add_heading('Risks & Issues', 2)
    risks_heading.runs[0].font.color.rgb = RGBColor(224, 112, 32)
    
    risks_table = doc.add_table(rows=3, cols=3)
    risks_table.style = 'Table Grid'
    
    # Headers
    headers = ['Type', 'Description', 'Mitigation / Action']
    for i, header in enumerate(headers):
        cell = risks_table.rows[0].cells[i]
        cell.text = header
        set_cell_color(cell, '262626')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Issues row
    risks_table.rows[1].cells[0].text = 'Issues'
    risks_table.rows[1].cells[1].text = str(project['issues'])
    risks_table.rows[1].cells[2].text = '[To be defined]'
    
    # Risks row
    risks_table.rows[2].cells[0].text = 'Risks'
    risks_table.rows[2].cells[1].text = str(project['risks'])
    risks_table.rows[2].cells[2].text = '[To be defined]'
    
    # Set gray background for data rows
    for row_idx in [1, 2]:
        for cell in risks_table.rows[row_idx].cells:
            set_cell_color(cell, 'F0F0F0')
    
    # Comments section if present
    if project['comments'] and project['comments'] != '[To be provided]':
        doc.add_paragraph()
        comments_heading = doc.add_heading('Comments / Notes', 2)
        comments_heading.runs[0].font.color.rgb = RGBColor(224, 112, 32)
        doc.add_paragraph(str(project['comments']))
    
    # Deliverables placeholder
    doc.add_paragraph()
    deliverables_heading = doc.add_heading('Deliverables / Milestones', 2)
    deliverables_heading.runs[0].font.color.rgb = RGBColor(224, 112, 32)
    
    deliverables_table = doc.add_table(rows=2, cols=5)
    deliverables_table.style = 'Table Grid'
    
    # Headers
    headers = ['Deliverable Name', 'Contractual Date', 'Planned Date', '% Done', 'Status']
    for i, header in enumerate(headers):
        cell = deliverables_table.rows[0].cells[i]
        cell.text = header
        set_cell_color(cell, '262626')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Placeholder row
    placeholder_values = ['[To be added manually]', '-', '-', '-', '-']
    for i, value in enumerate(placeholder_values):
        cell = deliverables_table.rows[1].cells[i]
        cell.text = value
        set_cell_color(cell, 'F0F0F0')


def generate_individual_word_report(project, output_path):
    """Generate a Word document for a single project."""
    return create_word_report([project], output_path)
